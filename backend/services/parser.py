import re
from urllib.parse import urlparse

from pypdf import PdfReader
import docx
from docx.oxml.ns import qn


def _normalize_text(text):
    return re.sub(r"\s+", " ", text.strip())


def _normalize_url(url):
    url = url.rstrip('.,;:')
    if not url.startswith(('http://', 'https://')):
        url = 'https://' + url
    return url


def _extract_links(text):
    links = {}
    url_pattern = r"((?:https?://)?(?:www\.)?(?:linkedin\.com/in/[A-Za-z0-9_-]+|github\.com/[A-Za-z0-9_-]+|medium\.com/[A-Za-z0-9_-]+|[A-Za-z0-9-]+\.(?:com|io|dev|me|tech|org)(?:/[A-Za-z0-9_\-./?=&%]+)?))"
    matches = re.findall(url_pattern, text)
    for match in matches:
        url = _normalize_url(match)
        parsed = urlparse(url)
        host = parsed.netloc.lower()
        if 'linkedin.com' in host:
            links['linkedin'] = url
        elif 'github.com' in host:
            links['github'] = url
        elif 'portfolio' in host or 'medium.com' in host:
            links['portfolio'] = url
        else:
            links.setdefault('website', url)
    return links


def _extract_docx_hyperlinks(doc):
    hyperlinks = []
    for paragraph in doc.paragraphs:
        for hyperlink in paragraph._p.findall('.//w:hyperlink', namespaces=paragraph._p.nsmap):
            rid = hyperlink.get(qn('r:id'))
            if not rid:
                continue
            target = paragraph.part.rels[rid].target_ref
            text = ''.join(node.text or '' for node in hyperlink.findall('.//w:t'))
            if target and text:
                hyperlinks.append((text, target))
    return hyperlinks


def _extract_emails(text):
    return re.findall(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}", text)


def _extract_phone(text):
    phones = re.findall(r"(?:\+\d{1,2}[\s-])?(?:\(?\d{3}\)?[\s.-]?\d{3}[\s.-]?\d{4})", text)
    return phones[:3]


def _extract_sections(text):
    sections = {
        'education': [],
        'experience': [],
        'projects': [],
        'skills': [],
        'achievements': [],
        'certifications': [],
        'internships': [],
        'leadership': [],
        'languages': []
    }

    lines = [line.strip() for line in text.splitlines() if line.strip()]
    for line in lines:
        lowered = line.lower()
        if 'education' in lowered or 'degree' in lowered or 'university' in lowered or 'college' in lowered:
            sections['education'].append(line)
        elif 'experience' in lowered or 'employment' in lowered or 'software engineer' in lowered or 'intern' in lowered:
            sections['experience'].append(line)
        elif 'project' in lowered:
            sections['projects'].append(line)
        elif 'skill' in lowered:
            sections['skills'].append(line)
        elif 'achievement' in lowered or 'award' in lowered:
            sections['achievements'].append(line)
        elif 'certification' in lowered or 'certificate' in lowered:
            sections['certifications'].append(line)
        elif 'language' in lowered:
            sections['languages'].append(line)

    return sections


def extract_resume_profile(resume_text):
    normalized = _normalize_text(resume_text)
    lines = [line.strip() for line in resume_text.splitlines() if line.strip()]

    name = lines[0] if lines else 'Unknown'
    emails = _extract_emails(normalized)
    phones = _extract_phone(normalized)
    links = _extract_links(resume_text)
    sections = _extract_sections(resume_text)

    skills = []
    for line in lines:
        if any(token in line.lower() for token in ['python', 'java', 'react', 'sql', 'docker', 'aws', 'kubernetes', 'javascript', 'node', 'c++', 'c#', 'machine learning', 'deep learning']):
            skills.extend([part.strip() for part in re.split(r'[,;|]', line) if part.strip()])
    skills = [skill.lower() for skill in skills if len(skill) > 2][:20]

    return {
        'name': name,
        'email': emails[0] if emails else '',
        'phone': phones[0] if phones else '',
        'location': '',
        'links': links,
        'skills': skills,
        'projects': [line for line in lines if 'project' in line.lower()][:5],
        'experience': [line for line in lines if 'experience' in line.lower() or 'intern' in line.lower()][:5],
        'education': sections['education'][:5],
        'sections': sections,
        'raw_text': resume_text
    }


def extract_text_from_pdf(pdf_path):
    text = ""
    try:
        reader = PdfReader(pdf_path)
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    except Exception as e:
        print(f"Error reading PDF {pdf_path}: {e}")
    return text


def extract_text_from_docx(docx_path):
    try:
        document = docx.Document(docx_path)
        paragraphs = []
        for paragraph in document.paragraphs:
            text = paragraph.text
            for hyperlink_text, hyperlink_target in _extract_docx_hyperlinks(document):
                if hyperlink_text and hyperlink_target and hyperlink_text in text and hyperlink_target not in text:
                    text = text.replace(hyperlink_text, f"{hyperlink_text} {hyperlink_target}")
            paragraphs.append(text)
        return "\n".join(paragraphs)
    except Exception as e:
        print(f"Error reading DOCX {docx_path}: {e}")
        return ""


def extract_text(filepath):
    lower_path = filepath.lower()
    if lower_path.endswith('.pdf'):
        return extract_text_from_pdf(filepath)
    elif lower_path.endswith('.docx'):
        return extract_text_from_docx(filepath)
    return ""